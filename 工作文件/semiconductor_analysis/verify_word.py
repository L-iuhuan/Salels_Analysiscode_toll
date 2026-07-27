from docx import Document
doc = Document(r"C:\Users\45091\Desktop\2026H1销售分析报告_深度版.docx")

print(f"段落数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")
print(f"节数: {len(doc.sections)}")

# Check headings
headings = []
for p in doc.paragraphs:
    if p.style.name.startswith('Heading'):
        headings.append((p.style.name, p.text[:80]))

print("\n== 标题结构 ==")
for style, text in headings:
    indent = {'Heading 1': '', 'Heading 2': '  ', 'Heading 3': '    '}.get(style, '')
    print(f"{indent}[{style}] {text}")

# Content checks
all_text = ' '.join([p.text for p in doc.paragraphs])

checks = [
    ('一页纸总览', ['34.3%', '246万', '核心悖论', '降价冲量']),
    ('一·整体态势', ['43,080', '14,832', '34.43%', '5,342', '交互残差', '3,024']),
    ('二·增长结构', ['通用电源管理', '有刷直流电机驱动', '62.14%', 'STI3452HFI', '产品力差']),
    ('三·新品引擎', ['5,241万', '12.2%', '39.21%', '89.7%']),
    ('四·毛利桥', ['价效应', '量效应', '结构效应', '成本效应', '交互残差', '可比SKU']),
    ('五·头号拖累', ['DCDC-18V', '中兴康讯', '共进', 'STI3452HFI']),
    ('六·品类总结', ['目标毛利率', '四档', '明星', '拖累']),
    ('七·负毛利', ['22个', '185.2', 'STI3452HFI']),
    ('八·大客户', ['追觅', '中兴康讯', '共进', '海康威视', '大客户板块']),
    ('九·销售需求', ['贺淼淼', '周小力', '刘仲涵', '颜蓉蓉', '袁坤']),
    ('十·价格ASP', ['含赠ASP', '付费ASP', '平均ASP', 'TMI8180G', '490.9']),
    ('十一·35%路径', ['四杠杆', '四条路径', '方案C', '1,208', '492%']),
    ('十二·决策KPI', ['限价', '停售', '挽回', '买赠专项审查']),
    ('附件', ['40个', 'KA/AA', '25人', '76品类', '四档']),
    ('总结', ['要决策', '要改变', '要建议', '止跌', '35%']),
]

print("\n== 内容覆盖检查 ==")
all_ok = True
for section, keywords in checks:
    found = all(k in all_text for k in keywords)
    status = '✓' if found else '✗ MISSING'
    if not found:
        all_ok = False
        missing = [k for k in keywords if k not in all_text]
        print(f"  {section}: {status}  missing: {missing}")
    else:
        print(f"  {section}: {status}")

print(f"\n结论: {'全部通过' if all_ok else '有遗漏'}")

# Count tables per section
print(f"\n表格数: {len(doc.tables)}")
