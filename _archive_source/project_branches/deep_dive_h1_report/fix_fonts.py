"""Word: 按字体方案统一设置所有文字,只改字体属性不动内容"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = r"C:\Users\45091\Desktop\2026H1销售分析报告_分析+附件.docx"
doc = Document(DOC_PATH)

INK = RGBColor(0x23,0x30,0x3F)
BLUE = RGBColor(0x14,0x3D,0x63)
GREEN = RGBColor(0x1F,0x7A,0x45)
RED = RGBColor(0xA8,0x23,0x0F)
GRAY = RGBColor(0x5D,0x6B,0x7A)
AMBER = RGBColor(0xB0,0x68,0x08)
WHITE = RGBColor(0xFF,0xFF,0xFF)

# ===== 1. 修改文档级样式 =====
for style_name, size, bold, color in [
    ('Heading 1', Pt(20), True, BLUE),
    ('Heading 2', Pt(14), True, BLUE),
    ('Heading 3', Pt(12), True, BLUE),
]:
    if style_name in [s.name for s in doc.styles]:
        st = doc.styles[style_name]
        st.font.size = size; st.font.bold = bold; st.font.color.rgb = color

# Normal style
ns = doc.styles['Normal']
ns.font.size = Pt(10); ns.font.bold = False; ns.font.color.rgb = INK

# ===== 2. 段落字体 =====
for p in doc.paragraphs:
    style_name = p.style.name if p.style else ''

    # 跳过 heading(由样式控制)
    if style_name.startswith('Heading'):
        continue

    # 检测callout(有段落底纹)
    is_callout = False
    pPr = p._p.find(qn('w:pPr'))
    if pPr is not None:
        shd = pPr.find(qn('w:shd'))
        if shd is not None:
            fill = shd.get(qn('w:fill'))
            if fill and fill not in ('auto','FFFFFF','none'):
                is_callout = True

    # 检测是否是KPI tile内的段落
    is_kpi_val = False; is_kpi_lbl = False
    if not is_callout and not style_name.startswith('Heading'):
        # 短文本,可能是tile
        pass

    for run in p.runs:
        rPr = run._r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr'); run._r.insert(0, rPr)

        if is_callout:
            run.font.size = Pt(9)
            run.font.color.rgb = INK
            # 首句加粗
            if run is p.runs[0]:
                run.font.bold = True
            else:
                run.font.bold = False
        else:
            run.font.size = Pt(10)
            run.font.color.rgb = INK
            run.font.bold = False

# ===== 3. 表格字体 =====
for tbl in doc.tables:
    if len(tbl.rows) == 0: continue

    # 判断首行是否为表头(有灰蓝背景)
    first_row = tbl.rows[0]
    is_header = False
    for cell in first_row.cells:
        tcPr = cell._tc.find(qn('w:tcPr'))
        if tcPr is not None:
            shd = tcPr.find(qn('w:shd'))
            if shd is not None:
                fill = shd.get(qn('w:fill'))
                if fill and fill.upper() == 'EEF2F8':
                    is_header = True
                    break

    for ri, row in enumerate(tbl.rows):
        for cell in row.cells:
            # 检测cell背景色
            tcPr = cell._tc.find(qn('w:tcPr'))
            cell_bg = None
            if tcPr is not None:
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    cell_bg = shd.get(qn('w:fill'))

            # 检测cell文本
            cell_text = cell.text.strip()

            for p in cell.paragraphs:
                for run in p.runs:
                    rPr = run._r.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = OxmlElement('w:rPr'); run._r.insert(0, rPr)

                    # 表头行
                    if ri == 0 and is_header:
                        run.font.size = Pt(8); run.font.bold = True; run.font.color.rgb = BLUE
                    else:
                        run.font.size = Pt(8); run.font.color.rgb = INK; run.font.bold = False

                    # 正负数字上色
                    t = run.text or ''
                    if t.startswith('+') and any(c.isdigit() for c in t):
                        run.font.color.rgb = GREEN; run.font.bold = True
                    elif t.startswith('−') or (t.startswith('−') and any(c.isdigit() for c in t)):
                        run.font.color.rgb = RED; run.font.bold = True
                    elif t and t[0] == '−' and any(c.isdigit() for c in t):
                        run.font.color.rgb = RED; run.font.bold = True
                    # 红/绿底单元格对应文字色
                    elif cell_bg and cell_bg.upper() == 'FBE6E1':
                        run.font.color.rgb = RED; run.font.bold = True
                    elif cell_bg and cell_bg.upper() == 'E7F3EC':
                        run.font.color.rgb = GREEN; run.font.bold = True
                    elif cell_bg and cell_bg.upper() == 'FBF0DA':
                        run.font.color.rgb = AMBER

# ===== 4. KPI tile区域特殊处理(基于文本模式) =====
for p in doc.paragraphs:
    if p.style.name.startswith('Heading'): continue
    text = p.text.strip()
    for run in p.runs:
        # 大数字tile: "+34.3%" "−2.37pct" "5,241万" 等短文本在表格外
        if len(text) < 20 and (text.endswith('%') or text.endswith('万') or text.endswith('pct')):
            if p.runs and run is p.runs[0]:
                # 检测是否在某个小表格里(tiles)
                pass

# ===== 5. 附件销售产品表(最后一组追加的):表头+GREEN_BG/RED_BG行处理 =====
# (已在上面table循环中处理)

doc.save(DOC_PATH)
print("DONE: 字体统一完成")
